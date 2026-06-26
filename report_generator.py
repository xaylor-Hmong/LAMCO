import os
import sys
import json
import math
import datetime as dt
from pathlib import Path
import requests

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Brand / style constants
# ----------------------------------------------------------------------------
NAVY = RGBColor(0x1F, 0x3A, 0x52)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE_FILL = "D9E8F5"
HEADER_GREY = "F2F2F2"
CHECK_ON = "\u2611"   # ☑
CHECK_OFF = "\u2610"  # ☐
FONT = "Phetsarath OT"

# ----------------------------------------------------------------------------
# Word document layout styling helpers
# ----------------------------------------------------------------------------
def _insert_in_order(tcPr, new_el, tag_order):
    new_tag = new_el.tag.split("}")[-1]
    if new_tag not in tag_order:
        tcPr.append(new_el)
        return
    new_idx = tag_order.index(new_tag)
    for existing in list(tcPr):
        existing_tag = existing.tag.split("}")[-1]
        if existing_tag in tag_order and tag_order.index(existing_tag) > new_idx:
            existing.addprevious(new_el)
            return
    tcPr.append(new_el)

_TCPR_ORDER = ["tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
               "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark"]

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _insert_in_order(tcPr, shd, _TCPR_ORDER)

def set_cell_borders(cell, color="BFBFBF", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    _insert_in_order(tcPr, borders, _TCPR_ORDER)

def style_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name = FONT
    
    # Force Complex Script for Lao language rendering
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), FONT)       
    rFonts.set(qn('w:eastAsia'), FONT) 
    
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

_PPR_ORDER = ["pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
              "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd",
              "tabs", "suppressAutoHyphens", "kinsoku", "wordWrap",
              "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
              "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
              "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc",
              "textDirection", "textAlignment", "textboxTightWrap",
              "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange"]

def _insert_pPr_in_order(pPr, new_el):
    new_tag = new_el.tag.split("}")[-1]
    if new_tag not in _PPR_ORDER:
        pPr.append(new_el)
        return
    new_idx = _PPR_ORDER.index(new_tag)
    for existing in list(pPr):
        existing_tag = existing.tag.split("}")[-1]
        if existing_tag in _PPR_ORDER and _PPR_ORDER.index(existing_tag) > new_idx:
            existing.addprevious(new_el)
            return
    pPr.append(new_el)

def add_bottom_rule(p, color="2E75B6", sz="12", space="4"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    _insert_pPr_in_order(pPr, pBdr)

def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    add_bottom_rule(p, color="2E75B6", sz="12", space="4")
    run = p.add_run(f"{number}.  {title}")
    style_run(run, size=13, bold=True, color=NAVY)
    return p

def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    style_run(run, size=11.5, bold=True, color=ACCENT_BLUE)
    return p

def add_label_value_row(table, label, value, label_width=Cm(4.2)):
    row = table.add_row()
    c0, c1 = row.cells
    c0.width = label_width
    set_cell_shading(c0, HEADER_GREY)
    set_cell_borders(c0)
    set_cell_borders(c1)
    p0 = c0.paragraphs[0]
    style_run(p0.add_run(label), size=10.5, bold=True)
    p1 = c1.paragraphs[0]
    style_run(p1.add_run(str(value) if value is not None and not (isinstance(value, float) and math.isnan(value)) else "\u2014"), size=10.5)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def checklist_line(doc, options_selected_pairs, indent=Cm(0.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(4)
    for i, (label, selected) in enumerate(options_selected_pairs):
        symbol = CHECK_ON if selected else CHECK_OFF
        run = p.add_run(f"{symbol} {label}")
        style_run(run, size=10.5, bold=selected)
        if i < len(options_selected_pairs) - 1:
            sep = p.add_run("      ")
            style_run(sep, size=10.5)
    return p

def add_caption_field(doc, label, value, indent=Cm(0.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(f"{label}: "), size=10.5, bold=True)
    style_run(p.add_run(str(value) if value is not None else "\u2014"), size=10.5)
    return p

def add_photo_block(doc, index, photo_path, fields, indent=Cm(0.5)):
    hdr = doc.add_paragraph()
    hdr.paragraph_format.left_indent = indent
    hdr.paragraph_format.space_before = Pt(10)
    style_run(hdr.add_run(f"Photo {index}"), size=10.5, bold=True, italic=True)

    if photo_path and Path(photo_path).exists():
        img_p = doc.add_paragraph()
        img_p.paragraph_format.left_indent = indent
        run = img_p.add_run()
        try:
            run.add_picture(photo_path, width=Cm(9.5))
        except Exception:
            style_run(img_p.add_run("[Photo could not be loaded]"), size=10, italic=True)
    else:
        miss_p = doc.add_paragraph()
        miss_p.paragraph_format.left_indent = indent
        style_run(miss_p.add_run("[No photo attached]"), size=10, italic=True, color=RGBColor(0x99, 0x99, 0x99))

    for label, value in fields:
        add_caption_field(doc, label, value, indent=indent)

def add_page_number_footer(doc, generated_at=None, version=1):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if generated_at:
        version_note = f" (v{version})" if version and version > 1 else ""
        pre_run = p.add_run(f"Generated {generated_at}{version_note}  \u2014  Page ")
        style_run(pre_run, size=9, color=RGBColor(0x80, 0x80, 0x80))
    run = p.add_run()
    style_run(run, size=9, color=RGBColor(0x80, 0x80, 0x80))
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)

def make_two_col_table(doc):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.8)
    return table

# ----------------------------------------------------------------------------
# Document Construction Loop
# ----------------------------------------------------------------------------
def build_report(data, photo_dir, output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)

    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10.5)

    # Force Lao font into the primary Word Style for regular paragraphs
    rFonts_normal = doc.styles["Normal"].font._element.rPr.get_or_add_rFonts()
    rFonts_normal.set(qn('w:cs'), FONT)
    rFonts_normal.set(qn('w:eastAsia'), FONT)

    add_page_number_footer(doc, generated_at=data.get("_generated_at"), version=data.get("_version", 1))

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title_p.add_run("LAO CONSULTING GROUP LTD."), size=15, bold=True, color=NAVY)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(sub_p.add_run("Daily Site Report \u2013 LAMCO Construction"), size=12, color=ACCENT_BLUE, bold=True)

    rule_p = doc.add_paragraph()
    add_bottom_rule(rule_p, color="1F3A52", sz="18", space="1")

    proj_hdr = doc.add_paragraph()
    style_run(proj_hdr.add_run("PROJECT DETAILS"), size=11.5, bold=True, color=NAVY)
    proj_table = make_two_col_table(doc)
    add_label_value_row(proj_table, "Project Name", data.get("project_name"))
    add_label_value_row(proj_table, "LCG Project Number", data.get("lcg_project_number"))
    add_label_value_row(proj_table, "Location", data.get("project_location"))
    add_label_value_row(proj_table, "Contractor", data.get("contractor_name"))
    add_label_value_row(proj_table, "Contract Number", data.get("contract_number"))

    doc.add_paragraph()
    daily_table = make_two_col_table(doc)
    add_label_value_row(daily_table, "Date", data.get("report_date"))
    weather = data.get("weather", "")
    weather_display = {
        "sunny": "Sunny", "raining": "Raining", "part rain - part sun": "Part rain / part sunny"
    }.get(weather, weather or "\u2014")
    add_label_value_row(daily_table, "Weather", weather_display)
    add_label_value_row(daily_table, "Reported by", f"{data.get('reporter_name','\u2014')} ({data.get('reporter_role','\u2014')})")

    add_section_heading(doc, 1, "Daily Work Summary")
    add_caption_field(doc, "Description of work this day", data.get("work_description"), indent=Cm(0))
    status = data.get("progress_status", "")
    status_opts = [("On track", status == "on track"), ("Ahead", status == "ahead"), ("Delayed", status == "delayed")]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    style_run(p.add_run("Progress status overall: "), size=10.5, bold=True)
    checklist_line(doc, status_opts, indent=Cm(0))

    add_section_heading(doc, 2, "Progress Photos")
    progress_photos = data.get("progress_photos", [])
    if not progress_photos:
        style_run(doc.add_paragraph().add_run("No progress photos submitted."), italic=True, size=10)
    for i, entry in enumerate(progress_photos, start=1):
        add_photo_block(doc, i, entry.get("photo_path"), [
            ("Location", entry.get("location_note")),
            ("Caption", entry.get("caption")),
        ])

    add_section_heading(doc, 3, "Quality and Testing")
    add_subheading(doc, "Inspections")
    inspections = data.get("inspections", [])
    if not inspections:
        style_run(doc.add_paragraph().add_run("No inspections submitted."), italic=True, size=10)
    for i, entry in enumerate(inspections, start=1):
        add_photo_block(doc, i, entry.get("photo_path"), [
            ("Location", entry.get("location")),
            ("Drg. No. ref.", entry.get("drawing_ref")),
            ("Description / Caption", entry.get("caption")),
        ])
        checklist_line(doc, [("Pass", entry.get("result") == "pass"), ("Fail", entry.get("result") == "fail")])
        checklist_line(doc, [("Follow-up required? Yes", entry.get("followup") == "yes"), ("No", entry.get("followup") == "no")])
        if entry.get("followup") == "yes" and entry.get("followup_notes"):
            add_caption_field(doc, "Follow-up notes", entry.get("followup_notes"))

    add_subheading(doc, "Testing")
    testing = data.get("testing", [])
    if not testing:
        style_run(doc.add_paragraph().add_run("No testing submitted."), italic=True, size=10)
    for i, entry in enumerate(testing, start=1):
        add_photo_block(doc, i, entry.get("photo_path"), [
            ("Location", entry.get("location")),
            ("Testing", entry.get("test_type")),
            ("Description / Caption", entry.get("caption")),
        ])
        checklist_line(doc, [("Pass", entry.get("result") == "pass"), ("Fail", entry.get("result") == "fail")])
        checklist_line(doc, [("Follow-up required? Yes", entry.get("followup") == "yes"), ("No", entry.get("followup") == "no")])
        if entry.get("followup") == "yes" and entry.get("followup_notes"):
            add_caption_field(doc, "Follow-up notes", entry.get("followup_notes"))

    add_section_heading(doc, 4, "Materials Delivered on Site")
    materials = data.get("materials", [])
    if not materials:
        style_run(doc.add_paragraph().add_run("No materials submitted."), italic=True, size=10)
    for i, entry in enumerate(materials, start=1):
        add_photo_block(doc, i, entry.get("photo_path"), [])
        add_caption_field(doc, "Material", entry.get("materials_name"))
        q = entry.get("quality", "")
        checklist_line(doc, [("Good", q == "good"), ("Acceptable", q == "acceptable"), ("Defective", q == "defective")])
        add_caption_field(doc, "Description", entry.get("description"))

    add_section_heading(doc, 5, "Safety and Environment")
    safety = data.get("safety", [])
    if not safety:
        style_run(doc.add_paragraph().add_run("No safety/environment items submitted."), italic=True, size=10)
    for i, entry in enumerate(safety, start=1):
        add_photo_block(doc, i, entry.get("photo_path"), [
            ("Location", entry.get("location")),
            ("Description / Issue noted", entry.get("issue")),
        ])

    doc.save(output_path)

# ----------------------------------------------------------------------------
# Kobo API parsing engine
# ----------------------------------------------------------------------------
def resolve_photo_path(field_value, attachment_map):
    if not field_value:
        return None
    if field_value in attachment_map:
        return attachment_map[field_value]
        
    def clean_name(name):
        return "".join([c.lower() for c in str(name) if c.isalnum()])
        
    target_clean = clean_name(Path(field_value).name)
    for original_key, local_path in attachment_map.items():
        if clean_name(Path(original_key).name) == target_clean:
            return local_path
    return None

def extract_repeat_group(submission, group_name, field_map):
    entries = submission.get(group_name) or submission.get(f"grp_quality/{group_name}")
    if not entries and "grp_quality" in submission:
        grp = submission.get("grp_quality")
        if isinstance(grp, list) and len(grp) > 0:
            entries = grp[0].get(group_name) or grp[0].get(f"grp_quality/{group_name}")
        elif isinstance(grp, dict):
            entries = grp.get(group_name) or grp.get(f"grp_quality/{group_name}")
            
    if not entries:
        entries = []
    
    out = []
    for entry in entries:
        item = {}
        for out_key, kobo_field in field_map.items():
            val = entry.get(kobo_field) or entry.get(f"{group_name}/{kobo_field}") or entry.get(f"grp_quality/{group_name}/{kobo_field}")
            if val is None:
                for actual_key, actual_val in entry.items():
                    if actual_key.endswith(f"/{kobo_field}") or actual_key == kobo_field:
                        val = actual_val
                        break
            item[out_key] = val
        out.append(item)
    return out

def build_data_dict(submission, attachment_map):
    def g(key):
        return submission.get(key) or submission.get(f"grp_project/{key}") or submission.get(f"grp_daily/{key}") \
            or submission.get(f"grp_summary/{key}") or submission.get(f"grp_reporter/{key}")

    data = {
        "project_name": g("project_name"),
        "lcg_project_number": g("lcg_project_number"),
        "project_location": g("project_location"),
        "contractor_name": g("contractor_name"),
        "contract_number": g("contract_number"),
        "report_date": g("report_date"),
        "weather": g("weather"),
        "reporter_name": g("reporter_name"),
        "reporter_role": g("reporter_role"),
        "work_description": g("work_description"),
        "progress_status": g("progress_status"),
    }

    p_photos = extract_repeat_group(submission, "rep_progress_photos", {"photo": "progress_photo", "location_note": "progress_photo_location_note", "caption": "progress_photo_caption"})
    for e in p_photos: e["photo_path"] = resolve_photo_path(e.get("photo"), attachment_map)
    data["progress_photos"] = p_photos

    insps = extract_repeat_group(submission, "rep_inspections", {"photo": "inspection_photo", "location": "inspection_location", "drawing_ref": "inspection_drawing_ref", "caption": "inspection_caption", "result": "inspection_result", "followup": "inspection_followup", "followup_notes": "inspection_followup_notes"})
    for e in insps: e["photo_path"] = resolve_photo_path(e.get("photo"), attachment_map)
    data["inspections"] = insps

    tests = extract_repeat_group(submission, "rep_testing", {"photo": "testing_photo", "location": "testing_location", "test_type": "testing_type", "caption": "testing_caption", "result": "testing_result", "followup": "testing_followup", "followup_notes": "testing_followup_notes"})
    for e in tests: e["photo_path"] = resolve_photo_path(e.get("photo"), attachment_map)
    data["testing"] = tests

    mats = extract_repeat_group(submission, "rep_materials", {"photo": "materials_photo", "materials_name": "materials_name", "quality": "materials_quality", "description": "materials_description"})
    for e in mats: e["photo_path"] = resolve_photo_path(e.get("photo"), attachment_map)
    data["materials"] = mats

    safeties = extract_repeat_group(submission, "rep_safety", {"photo": "safety_photo", "location": "safety_location", "issue": "safety_issue"})
    for e in safeties: e["photo_path"] = resolve_photo_path(e.get("photo"), attachment_map)
    data["safety"] = safeties

    return data
